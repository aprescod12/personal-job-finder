import hashlib
import io
import shutil
import tempfile

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document
from pypdf import PdfWriter

from tracker.models import CareerProfile

from .models import ResumeExtractionReview, ResumeSource
from .services.resume_deterministic import DeterministicResumeExtractor
from .services.resume_documents import (
    ResumeDocumentError,
    extract_resume_document_text,
)
from .services.resume_extraction import (
    BaseResumeExtractor,
    ResumeExtractionError,
    ResumeExtractionRequest,
    execute_resume_extractor,
    get_resume_extractor,
)
from .views import RESUME_EXTRACTION_SESSION_KEY


SAMPLE_RESUME_TEXT = """Amiri Prescod
amiri@example.com | +1 610-555-0100 | github.com/aprescod12

Professional Summary
Electrical engineer pursuing graduate study in biomedical engineering.

Education
Villanova University
B.S. Electrical Engineering | 2025
M.S. Biomedical Engineering Candidate | 2027

Experience
Engineering Intern
Medical Device Company | Summer 2025
Developed and tested embedded sensing prototypes.

Projects
Wearable Sensing Platform
Python, signal processing, embedded systems
Built a prototype for physiological data collection.

Technical Skills
Python, MATLAB, C, Embedded systems, Test and validation

Leadership and Activities
Student-Athlete
Villanova Track and Field
"""


class GuardedAIResumeExtractor(BaseResumeExtractor):
    provider_key = "test-ai"
    provider_label = "Test AI resume extractor"
    provider_version = "test-v1"
    extraction_mode = "ai"
    requires_ai_enabled = True

    def extract(self, request):
        return self.result(identity={}, profile={})


class InvalidResumeExtractor(BaseResumeExtractor):
    def extract(self, request):
        return {"not": "a ResumeExtractionResult"}


class ResumeExtractionContractTests(TestCase):
    def _request(self, text=SAMPLE_RESUME_TEXT):
        return ResumeExtractionRequest(
            document_text=text,
            source_id=1,
            source_sha256="a" * 64,
            source_filename="resume.txt",
            source_label="Test resume",
            document_parser_key="plain-text",
            document_parser_version="test-reader-v1",
        )

    def test_request_rejects_blank_document_text(self):
        with self.assertRaisesRegex(
            ResumeExtractionError,
            "Readable resume text is required",
        ):
            self._request("   ")

    def test_deterministic_provider_returns_json_safe_structured_draft(self):
        result = DeterministicResumeExtractor().extract(self._request()).to_dict()

        self.assertEqual(result["provider"]["key"], "deterministic")
        self.assertEqual(result["identity"]["full_name"], "Amiri Prescod")
        self.assertEqual(result["identity"]["email"], "amiri@example.com")
        self.assertIn("Python", result["profile"]["skills"])
        self.assertTrue(result["profile"]["education"])
        self.assertTrue(result["profile"]["experience"])
        self.assertTrue(result["profile"]["projects"])
        self.assertTrue(result["evidence"])

    def test_execute_rejects_wrong_result_type(self):
        with self.assertRaisesRegex(
            ResumeExtractionError,
            "did not return a ResumeExtractionResult",
        ):
            execute_resume_extractor(self._request(), InvalidResumeExtractor())

    @override_settings(
        RESUME_EXTRACTOR=(
            "candidate_profile.test_resume_extraction.GuardedAIResumeExtractor"
        ),
        RESUME_AI_ENABLED=False,
    )
    def test_ai_provider_requires_separate_enable_switch(self):
        with self.assertRaisesRegex(
            ResumeExtractionError,
            "RESUME_AI_ENABLED safety switch",
        ):
            get_resume_extractor()


class ResumeDocumentReaderTests(TestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls.media_root = tempfile.mkdtemp(prefix="resume-reader-tests-")
        cls.media_override = override_settings(MEDIA_ROOT=cls.media_root)
        cls.media_override.enable()

    @classmethod
    def tearDownClass(cls):
        cls.media_override.disable()
        shutil.rmtree(cls.media_root, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.profile = CareerProfile.get_solo()

    def _source(self, filename, content, content_type):
        return ResumeSource.objects.create(
            profile=self.profile,
            document=SimpleUploadedFile(filename, content, content_type=content_type),
            original_filename=filename,
            content_type=content_type,
            file_size=len(content),
            sha256=hashlib.sha256(content).hexdigest(),
            is_active=True,
        )

    def test_plain_text_reader_preserves_resume_content(self):
        content = SAMPLE_RESUME_TEXT.encode("utf-8")
        source = self._source("resume.txt", content, "text/plain")

        result = extract_resume_document_text(source)

        self.assertEqual(result.parser_key, "plain-text")
        self.assertIn("Amiri Prescod", result.text)
        self.assertIn("Technical Skills", result.text)

    def test_docx_reader_extracts_paragraphs_and_table_cells(self):
        buffer = io.BytesIO()
        document = Document()
        document.add_paragraph("Amiri Prescod")
        document.add_heading("Education", level=1)
        document.add_paragraph("Villanova University")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Python"
        table.cell(0, 1).text = "MATLAB"
        document.save(buffer)
        content = buffer.getvalue()
        source = self._source(
            "resume.docx",
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        result = extract_resume_document_text(source)

        self.assertEqual(result.parser_key, "python-docx")
        self.assertIn("Villanova University", result.text)
        self.assertIn("Python | MATLAB", result.text)

    def test_pdf_without_selectable_text_returns_clear_error(self):
        buffer = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        writer.write(buffer)
        content = buffer.getvalue()
        source = self._source("blank.pdf", content, "application/pdf")

        with self.assertRaisesRegex(ResumeDocumentError, "No readable text was found"):
            extract_resume_document_text(source)


@override_settings(
    RESUME_EXTRACTOR=(
        "candidate_profile.services.resume_deterministic.DeterministicResumeExtractor"
    )
)
class ResumeExtractionWorkflowTests(TestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls.media_root = tempfile.mkdtemp(prefix="resume-extraction-workflow-tests-")
        cls.media_override = override_settings(MEDIA_ROOT=cls.media_root)
        cls.media_override.enable()

    @classmethod
    def tearDownClass(cls):
        cls.media_override.disable()
        shutil.rmtree(cls.media_root, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.profile = CareerProfile.get_solo()
        self.profile.professional_headline = "Approved profile headline"
        self.profile.skills = "Approved skill one\nApproved skill two"
        self.profile.save()
        content = SAMPLE_RESUME_TEXT.encode("utf-8")
        self.source = ResumeSource.objects.create(
            profile=self.profile,
            document=SimpleUploadedFile(
                "amiri-resume.txt",
                content,
                content_type="text/plain",
            ),
            original_filename="amiri-resume.txt",
            label="Stage 5 test resume",
            content_type="text/plain",
            file_size=len(content),
            sha256=hashlib.sha256(content).hexdigest(),
            is_active=True,
        )

    def _extract(self):
        return self.client.post(
            reverse(
                "candidate_profile:run_resume_extraction",
                args=[self.source.id],
            )
        )

    def test_extraction_endpoint_requires_post(self):
        response = self.client.get(
            reverse(
                "candidate_profile:run_resume_extraction",
                args=[self.source.id],
            )
        )

        self.assertEqual(response.status_code, 405)
        self.assertNotIn(RESUME_EXTRACTION_SESSION_KEY, self.client.session)

    def test_extraction_creates_persistent_review_without_changing_profile(self):
        profile_before = {
            "professional_headline": self.profile.professional_headline,
            "skills": self.profile.skills,
            "updated_at": self.profile.updated_at,
        }

        response = self._extract()

        self.assertRedirects(
            response,
            reverse("candidate_profile:resume_extraction_review"),
        )
        session_value = self.client.session[RESUME_EXTRACTION_SESSION_KEY]
        review = ResumeExtractionReview.objects.get(id=session_value["review_id"])
        self.assertEqual(review.source_id, self.source.id)
        self.assertEqual(review.source_sha256, self.source.sha256)
        self.assertEqual(
            review.claims.get(claim_key="identity.full_name").reviewed_value,
            "Amiri Prescod",
        )
        self.assertTrue(review.claims.exists())

        self.profile.refresh_from_db()
        self.assertEqual(
            self.profile.professional_headline,
            profile_before["professional_headline"],
        )
        self.assertEqual(self.profile.skills, profile_before["skills"])
        self.assertEqual(self.profile.updated_at, profile_before["updated_at"])

    def test_review_page_discloses_controlled_persistence_boundary(self):
        self._extract()

        response = self.client.get(
            reverse("candidate_profile:resume_extraction_review")
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "MANUAL PROFILE UNCHANGED")
        self.assertContains(response, "Deterministic local resume parser")
        self.assertContains(response, "Amiri Prescod")
        self.assertContains(response, "SAVE REVIEW ONLY")
        self.assertContains(response, "APPLY APPROVED CLAIMS")
        self.assertNotContains(response, "APPLY TO PROFILE")

    def test_close_marks_review_discarded_and_keeps_resume_source(self):
        self._extract()
        review_id = self.client.session[RESUME_EXTRACTION_SESSION_KEY]["review_id"]

        response = self.client.post(
            reverse("candidate_profile:clear_resume_extraction")
        )

        self.assertRedirects(
            response,
            reverse("candidate_profile:resume_source_list"),
        )
        self.assertNotIn(RESUME_EXTRACTION_SESSION_KEY, self.client.session)
        self.assertTrue(ResumeSource.objects.filter(id=self.source.id).exists())
        review = ResumeExtractionReview.objects.get(id=review_id)
        self.assertEqual(review.status, ResumeExtractionReview.Status.DISCARDED)

    def test_unreadable_source_does_not_create_review(self):
        self.source.document.delete(save=False)
        blank_content = b"   \n\n"
        self.source.document = SimpleUploadedFile(
            "blank.txt",
            blank_content,
            content_type="text/plain",
        )
        self.source.original_filename = "blank.txt"
        self.source.file_size = len(blank_content)
        self.source.sha256 = hashlib.sha256(blank_content).hexdigest()
        self.source.save()

        response = self._extract()

        self.assertRedirects(
            response,
            reverse("candidate_profile:resume_source_list"),
        )
        self.assertNotIn(RESUME_EXTRACTION_SESSION_KEY, self.client.session)
        self.assertFalse(ResumeExtractionReview.objects.exists())
