import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from pypdf import PdfReader

if TYPE_CHECKING:
    from candidate_profile.models import ResumeSource


DOCUMENT_PARSER_VERSION = "local-resume-document-reader-v2"
MAX_EXTRACTED_CHARACTERS = 120_000
PDF_LAYOUT_SCALE_WEIGHT = 0.5

ERROR_UNSUPPORTED_FORMAT = "unsupported_format"
ERROR_UNREADABLE_DOCUMENT = "unreadable_document"
ERROR_EMPTY_DOCUMENT = "empty_document"

_BULLET_TRANSLATION = str.maketrans(
    {
        "▪": "•",
        "◦": "•",
        "●": "•",
        "·": "•",
        "‣": "•",
    }
)


class ResumeDocumentError(RuntimeError):
    """Raised when a stored resume cannot be converted into readable text."""

    def __init__(self, message: str, *, category: str = ERROR_UNREADABLE_DOCUMENT):
        super().__init__(message)
        self.category = category


@dataclass(slots=True)
class ResumeDocumentText:
    text: str
    parser_key: str
    parser_version: str = DOCUMENT_PARSER_VERSION
    warnings: list[str] = field(default_factory=list)

    def to_dict(self):
        return {
            "text": self.text,
            "parser_key": self.parser_key,
            "parser_version": self.parser_version,
            "warnings": list(self.warnings),
        }


def _remove_hidden_format_characters(value: str) -> str:
    return "".join(
        character
        for character in value
        if character in {"\n", "\t"}
        or unicodedata.category(character) != "Cf"
    )


def _normalize_line(value: str) -> str:
    value = value.translate(_BULLET_TRANSLATION)
    value = re.sub(r"[ \t]+", " ", value).strip()
    value = re.sub(r"\s*[_=]{3,}\s*$", "", value).rstrip()
    if value.startswith("•"):
        value = f"• {value[1:].lstrip()}"
    return value


def _normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFKC", (value or "").replace("\u00a0", " "))
    value = _remove_hidden_format_characters(value)
    value = value.replace("\r\n", "\n").replace("\r", "\n")

    output_lines: list[str] = []
    continuing_bullet = False
    for raw_line in value.split("\n"):
        leading_spaces = len(raw_line) - len(raw_line.lstrip(" \t"))
        line = _normalize_line(raw_line)
        if not line:
            if output_lines and output_lines[-1] != "":
                output_lines.append("")
            continuing_bullet = False
            continue

        is_bullet = line.startswith("• ")
        is_indented_continuation = leading_spaces >= 5 and not is_bullet
        if output_lines and output_lines[-1]:
            previous = output_lines[-1]
            if previous.endswith("-") and line[:1].islower():
                output_lines[-1] = f"{previous}{line}"
                continuing_bullet = previous.startswith("• ") or continuing_bullet
                continue
            if continuing_bullet and is_indented_continuation:
                output_lines[-1] = f"{previous} {line}"
                continue

        output_lines.append(line)
        continuing_bullet = is_bullet

    value = "\n".join(output_lines)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _read_txt(stream) -> tuple[str, list[str]]:
    raw_bytes = stream.read()
    warnings = []
    try:
        text = raw_bytes.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = raw_bytes.decode("utf-8", errors="replace")
        warnings.append(
            "Some plain-text characters could not be decoded exactly and were replaced."
        )
    return text, warnings


def _extract_pdf_page_text(page) -> tuple[str, bool]:
    try:
        text = page.extract_text(
            extraction_mode="layout",
            layout_mode_scale_weight=PDF_LAYOUT_SCALE_WEIGHT,
        )
        return text or "", True
    except TypeError:
        return page.extract_text() or "", False


def _read_pdf(stream) -> tuple[str, list[str]]:
    warnings: list[str] = []
    layout_mode_used = True
    try:
        reader = PdfReader(stream)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:  # pragma: no cover - provider-specific failure
                raise ResumeDocumentError(
                    "The PDF is password protected and could not be read."
                ) from exc

        pages = []
        empty_pages = []
        for page_number, page in enumerate(reader.pages, start=1):
            page_text, used_layout = _extract_pdf_page_text(page)
            layout_mode_used = layout_mode_used and used_layout
            if page_text.strip():
                pages.append(page_text)
            else:
                empty_pages.append(page_number)
    except ResumeDocumentError:
        raise
    except Exception as exc:
        raise ResumeDocumentError(
            "The stored PDF could not be read safely."
        ) from exc

    if not layout_mode_used:
        warnings.append(
            "The installed PDF reader did not support layout-preserving extraction; plain extraction was used."
        )
    if empty_pages:
        page_list = ", ".join(str(number) for number in empty_pages[:8])
        suffix = "" if len(empty_pages) <= 8 else ", …"
        warnings.append(
            f"No selectable text was found on PDF page(s): {page_list}{suffix}."
        )
    return "\n\n".join(pages), warnings


def _read_docx(stream) -> tuple[str, list[str]]:
    try:
        document = Document(stream)
    except Exception as exc:
        raise ResumeDocumentError(
            "The stored DOCX file could not be read safely."
        ) from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cell_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_values:
                parts.append(" | ".join(cell_values))
    return "\n".join(parts), []


def extract_resume_document_text(source: "ResumeSource") -> ResumeDocumentText:
    """Read a stored resume into text without changing any model records."""

    extension = Path(source.original_filename or source.document.name).suffix.casefold()
    parser_map = {
        ".txt": ("plain-text", _read_txt),
        ".pdf": ("pypdf", _read_pdf),
        ".docx": ("python-docx", _read_docx),
    }
    parser = parser_map.get(extension)
    if parser is None:
        raise ResumeDocumentError(
            "This resume format does not have a configured document reader.",
            category=ERROR_UNSUPPORTED_FORMAT,
        )

    parser_key, reader = parser
    try:
        with source.document.open("rb") as stream:
            text, warnings = reader(stream)
    except ResumeDocumentError:
        raise
    except (OSError, ValueError) as exc:
        raise ResumeDocumentError(
            "The stored resume file could not be opened."
        ) from exc

    text = _normalize_text(text)
    if not text:
        raise ResumeDocumentError(
            "No readable text was found in this resume. Scanned PDFs require OCR, which is not enabled.",
            category=ERROR_EMPTY_DOCUMENT,
        )

    if len(text) > MAX_EXTRACTED_CHARACTERS:
        text = text[:MAX_EXTRACTED_CHARACTERS].rstrip()
        warnings.append(
            "The extracted document text exceeded the safety limit and was truncated."
        )

    return ResumeDocumentText(
        text=text,
        parser_key=parser_key,
        warnings=warnings,
    )
